from __future__ import annotations

import argparse
import json
import time
from pathlib import Path

from core.job_engine import init_progress, update_progress
from modules.gemini_client import GeminiClient

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _read_json(p: Path) -> dict:
    return json.loads(p.read_text(encoding="utf-8"))


def _log(log_path: Path, s: str):
    log_path.parent.mkdir(parents=True, exist_ok=True)
    with log_path.open("a", encoding="utf-8") as f:
        f.write(s.rstrip() + "\n")


def _apply_docx_format(doc: Document, fmt: dict):
    # A4 + margins 4-3-3-3 cm
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    m = fmt.get("margins_cm") or {}
    sec.top_margin = Cm(float(m.get("top", 4.0)))
    sec.right_margin = Cm(float(m.get("right", 3.0)))
    sec.bottom_margin = Cm(float(m.get("bottom", 3.0)))
    sec.left_margin = Cm(float(m.get("left", 3.0)))

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = fmt.get("font", "Times New Roman")
    style.font.size = Pt(int(fmt.get("font_size_pt", 12)))

    # line spacing 1.5 for normal paragraphs (we set per paragraph as we add)


def _add_para(doc: Document, text: str, *, align=None, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bool(bold)
    p.paragraph_format.line_spacing = 1.5
    if align is not None:
        p.alignment = align
    return p


def _prompt_section(lang: str, mode: str, meta: dict, sec_title: str, topic: str, goals: str, refs: list[str], target_pages: int) -> str:
    lang_name = "Bahasa Indonesia akademik formal" if lang == "id" else "Formal academic English"
    refs_txt = "\n".join(f"- {r}" for r in refs) if refs else "- (Tidak ada referensi wajib)"
    goals_txt = goals.strip() if goals.strip() else "(Tidak ada tambahan rumusan masalah khusus)"

    return f"""
Tulis bagian karya tulis akademik dengan gaya: {lang_name}.
Mode dokumen: {mode}.
Judul: {meta.get('title','')}

Bagian yang harus ditulis sekarang: "{sec_title}"

Konteks topik:
{topic}

Rumusan masalah/tujuan (jika ada):
{goals_txt}

Referensi wajib (jika relevan, gunakan sebagai rujukan konseptual, jangan halusin detail bibliografi):
{refs_txt}

Aturan:
- Jangan menulis daftar isi otomatis.
- Gunakan paragraf akademik, koheren, tidak bertele-tele.
- Hindari klaim data spesifik jika tidak ada rujukan.
- Panjang bagian menyesuaikan dokumen total kira-kira {target_pages} halaman A4 (1.5 spasi) — buat proporsional.
Output: hanya isi teks untuk bagian ini (tanpa judul ulang).
""".strip()


def main(cfg_path: str) -> int:
    cfgp = Path(cfg_path).resolve()
    cfg = _read_json(cfgp)

    job_dir = Path(cfg.get("job_dir") or cfgp.parent).resolve()
    log_path = Path(cfg.get("log_path") or (job_dir / "job.log")).resolve()

    api_key = (cfg.get("api_key") or "").strip()
    if not api_key:
        _log(log_path, "[ERROR] missing api_key")
        update_progress(job_dir, status="error", total=1, done=1, current="missing api_key")
        return 2

    mode = cfg["mode"]
    lang = cfg["lang"]
    template = cfg["template"]
    rename_map = cfg["rename_map"]
    meta = cfg["meta"]
    topic = cfg.get("topic", "")
    goals = cfg.get("goals", "")
    refs = cfg.get("refs") or []
    fmt = cfg.get("format") or {}
    target_pages = int(cfg.get("target_pages") or 20)

    client = GeminiClient(api_key=api_key)

    out_dir = job_dir / "outputs"
    out_dir.mkdir(parents=True, exist_ok=True)

    total = max(1, len(template))
    init_progress(job_dir, total=total)

    doc = Document()
    _apply_docx_format(doc, fmt)

    _log(log_path, f"[START] mode={mode} lang={lang} sections={len(template)}")

    for i, sec in enumerate(template, start=1):
        title = rename_map.get(sec["id"], sec["title"]).strip() or sec["title"]
        level = int(sec.get("level") or 0)

        update_progress(job_dir, status="running", total=total, done=i - 1, current=f"Writing: {title}")
        _log(log_path, f"[SECTION] {i}/{total} {title}")

        # headings
        if level == 0:
            _add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        else:
            # Heading styles (still keep spacing 1.5)
            p = doc.add_paragraph(title)
            p.style = "Heading 1" if level == 1 else "Heading 2"
            p.paragraph_format.line_spacing = 1.5

        # special sections (cover) minimal, else generate content
        sid = sec["id"]
        if sid in ("cover", "judul"):
            _add_para(doc, meta.get("title", ""), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
            if meta.get("author"):
                _add_para(doc, meta["author"], align=WD_ALIGN_PARAGRAPH.CENTER)
            if meta.get("institution"):
                _add_para(doc, meta["institution"], align=WD_ALIGN_PARAGRAPH.CENTER)
            if meta.get("program"):
                _add_para(doc, meta["program"], align=WD_ALIGN_PARAGRAPH.CENTER)
            if meta.get("year"):
                _add_para(doc, meta["year"], align=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_page_break()
            update_progress(job_dir, status="running", total=total, done=i, current=f"Done: {title}")
            continue

        if sid in ("daftar_pustaka",):
            # simple placeholder — can be improved to APA generator later
            if refs:
                for r in refs:
                    _add_para(doc, r)
            else:
                _add_para(doc, "Referensi disusun sesuai gaya sitasi institusi (mis. APA/IEEE) dan dilengkapi manual.")
            update_progress(job_dir, status="running", total=total, done=i, current=f"Done: {title}")
            continue

        prompt = _prompt_section(lang, mode, meta, title, topic, goals, refs, target_pages)
        resp = client.generate(prompt, temperature=0.4, max_tokens=1400, retries=3)
        text = (resp.get("text") or "").strip()

        # write paragraphs
        for para in [p.strip() for p in text.split("\n") if p.strip()]:
            _add_para(doc, para)

        doc.add_paragraph("")  # spacing
        update_progress(job_dir, status="running", total=total, done=i, current=f"Done: {title}")

    docx_path = out_dir / "karya_tulis.docx"
    doc.save(str(docx_path))

    update_progress(job_dir, status="done", total=total, done=total, current="done")
    _log(log_path, "[DONE] wrote DOCX")
    return 0


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", required=True)
    args = ap.parse_args()
    raise SystemExit(main(args.config))
