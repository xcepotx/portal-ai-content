# yt-automation-onefact-ind/tools/ebook_maker_worker.py
from __future__ import annotations

import argparse
import json
import os
import io
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional
from PIL import Image, ImageDraw, ImageFont
from modules.nano_banana_client import NanoBananaClient

import sys

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))

from core.job_engine import init_progress, update_progress  # noqa: E402


def _append_log(log_path: Path, msg: str):
    log_path.parent.mkdir(parents=True, exist_ok=True)
    line = f"[{time.strftime('%Y-%m-%d %H:%M:%S')}] {msg}\n"
    with log_path.open("a", encoding="utf-8") as f:
        f.write(line)


def _safe_slug(s: str, max_len: int = 40) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"\s+", "-", s)
    s = re.sub(r"[^a-z0-9\-]+", "-", s)
    s = re.sub(r"-{2,}", "-", s).strip("-")
    return (s or "chapter")[:max_len]


def _is_transient_error(e: Exception) -> bool:
    msg = str(e)
    name = type(e).__name__
    return (
        "RemoteProtocolError" in name
        or "Server disconnected" in msg
        or "429" in msg
        or "503" in msg
        or "RESOURCE_EXHAUSTED" in msg
        or "UNAVAILABLE" in msg
        or "No images returned" in msg
        or "timeout" in msg.lower()
    )


@dataclass
class RetryCfg:
    max_attempts: int = 6
    base_delay: float = 1.0
    max_delay: float = 20.0


def _genai_client(api_key: str):
    from google import genai  # local import
    return genai.Client(api_key=api_key, http_options={"timeout": 600000})


def _gen_text(
    *,
    api_key: str,
    model: str,
    prompt: str,
    log_path: Path,
    retry: RetryCfg,
    temperature: float = 0.7,
) -> str:
    from google.genai import types  # local import

    client = _genai_client(api_key)
    last_err: Exception | None = None

    for attempt in range(1, retry.max_attempts + 1):
        try:
            resp = client.models.generate_content(
                model=model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_modalities=["TEXT"],
                    temperature=temperature,
                ),
            )
            txt = (getattr(resp, "text", None) or "").strip()
            if not txt:
                # fallback parts
                parts = []
                for p in getattr(resp, "parts", []) or []:
                    if getattr(p, "text", None):
                        parts.append(p.text)
                txt = ("\n".join(parts)).strip()
            return txt
        except Exception as e:
            last_err = e
            _append_log(log_path, f"WARN: gen_text attempt {attempt}/{retry.max_attempts} failed: {type(e).__name__}: {e}")
            if not _is_transient_error(e):
                break
            delay = min(retry.max_delay, retry.base_delay * (1.6 ** (attempt - 1)))
            time.sleep(delay)

    raise RuntimeError(f"gen_text failed: {type(last_err).__name__}: {last_err}")


def _parse_outline(text: str, chapters_target: int) -> List[Dict[str, Any]]:
    """
    Terima outline dari model (bebas), coba parsing jadi list:
    [{"title": "...", "bullets": ["..",".."]}, ...]
    """
    text = (text or "").strip()

    # Try JSON first
    try:
        j = json.loads(text)
        if isinstance(j, dict) and "chapters" in j:
            j = j["chapters"]
        if isinstance(j, list):
            out = []
            for it in j:
                if isinstance(it, dict) and it.get("title"):
                    out.append({"title": str(it["title"]).strip(), "bullets": list(it.get("bullets") or [])})
                elif isinstance(it, str):
                    out.append({"title": it.strip(), "bullets": []})
            if out:
                return out[:chapters_target]
    except Exception:
        pass

    # Fallback: parse numbered lines
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    chapters: List[Dict[str, Any]] = []
    cur: Optional[Dict[str, Any]] = None
    for ln in lines:
        m = re.match(r"^\s*(\d+)[\)\.\-]\s*(.+)$", ln)
        if m:
            if cur:
                chapters.append(cur)
            cur = {"title": m.group(2).strip(), "bullets": []}
            continue
        if ln.startswith(("-", "*")) and cur is not None:
            cur["bullets"].append(ln.lstrip("-* ").strip())
    if cur:
        chapters.append(cur)

    # If still empty, make dummy
    if not chapters:
        chapters = [{"title": f"Chapter {i}", "bullets": []} for i in range(1, chapters_target + 1)]
    return chapters[:chapters_target]


def _outline_prompt(cfg: dict) -> str:
    title = (cfg.get("title") or "").strip()
    topic = (cfg.get("topic") or "").strip()
    language = (cfg.get("language") or "Indonesian").strip()
    tone = (cfg.get("tone") or "").strip()
    audience = (cfg.get("audience") or "").strip()
    chapters = int(cfg.get("chapters") or 8)

    return (
        "You are a professional book editor.\n"
        f"Language: {language}\n"
        f"Tone: {tone}\n"
        f"Target audience: {audience}\n"
        f"Book title: {title}\n"
        f"Brief/topic: {topic}\n\n"
        f"Create an outline with exactly {chapters} chapters.\n"
        "For each chapter provide:\n"
        "- title\n"
        "- 4-6 bullet points (what the chapter will cover)\n\n"
        "Return ONLY JSON with this shape:\n"
        "{\n"
        '  "chapters": [\n'
        '    {"title": "...", "bullets": ["...", "..."]}\n'
        "  ]\n"
        "}\n"
    )


def _chapter_prompt(cfg: dict, idx: int, ch: Dict[str, Any]) -> str:
    title = (cfg.get("title") or "").strip()
    language = (cfg.get("language") or "Indonesian").strip()
    tone = (cfg.get("tone") or "").strip()
    audience = (cfg.get("audience") or "").strip()
    words = int(cfg.get("words_per_chapter") or 650)

    bullets = ch.get("bullets") or []
    bullets_txt = "\n".join([f"- {b}" for b in bullets]) if bullets else "- (use your best judgement)"

    return (
        "You are a senior non-fiction writer.\n"
        f"Language: {language}\n"
        f"Tone: {tone}\n"
        f"Target audience: {audience}\n"
        f"Book title: {title}\n\n"
        f"Write Chapter {idx}: {ch.get('title','')}\n"
        f"Target length: ~{words} words.\n\n"
        "Must follow:\n"
        "- Use Markdown.\n"
        "- Start with '## Chapter {idx}: <chapter title>'\n"
        "- Use short sections with '###' headings.\n"
        "- Include bullet lists and at least 1 practical example or checklist.\n"
        "- Avoid repetition across chapters.\n\n"
        "Chapter coverage bullets:\n"
        f"{bullets_txt}\n"
    )


def _write_docx(md_text: str, out_path: Path, title: str, subtitle: str, author: str, cover_path: Optional[Path] = None):
    from docx import Document
    from docx.shared import Inches

    doc = Document()

    # ✅ COVER PAGE
    if cover_path and cover_path.exists():
        doc.add_picture(str(cover_path), width=Inches(6.0))
        doc.add_page_break()

    if title:
        doc.add_heading(title, level=0)
    if subtitle:
        doc.add_paragraph(subtitle)
    if author:
        doc.add_paragraph(f"By {author}")
    doc.add_page_break()

    lines = md_text.splitlines()
    for ln in lines:
        if ln.startswith("## "):
            doc.add_heading(ln[3:].strip(), level=1)
        elif ln.startswith("### "):
            doc.add_heading(ln[4:].strip(), level=2)
        elif ln.startswith("- "):
            doc.add_paragraph(ln[2:].strip(), style="List Bullet")
        elif ln.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(ln)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

def _write_pdf(md_text: str, out_path: Path, title: str, subtitle: str, author: str, cover_path: Optional[Path] = None):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    from PIL import Image as PILImage

    styles = getSampleStyleSheet()

    out_path.parent.mkdir(parents=True, exist_ok=True)

    # Normal pages (punya margin)
    doc = SimpleDocTemplate(str(out_path), pagesize=LETTER)

    def draw_cover(canvas, doc_):
        """Draw full-page cover on first page (no frame constraint)."""
        if not (cover_path and cover_path.exists()):
            return

        page_w, page_h = doc_.pagesize

        im = PILImage.open(cover_path)
        iw, ih = im.size  # pixels; reportlab treats px ~ points at 72dpi
        scale = min(page_w / float(iw), page_h / float(ih))
        w = iw * scale
        h = ih * scale
        x = (page_w - w) / 2.0
        y = (page_h - h) / 2.0

        canvas.drawImage(
            ImageReader(im),
            x, y,
            width=w, height=h,
            preserveAspectRatio=True,
            mask="auto",
        )

    story = []

    # Kalau ada cover: halaman 1 untuk cover (digambar via canvas), konten mulai halaman 2
    if cover_path and cover_path.exists():
        story.append(PageBreak())

    # Title/content
    if title:
        story.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
        story.append(Spacer(1, 0.15 * inch))
    if subtitle:
        story.append(Paragraph(subtitle, styles["Heading3"]))
    if author:
        story.append(Paragraph(f"By {author}", styles["Normal"]))
    story.append(Spacer(1, 0.35 * inch))

    for ln in md_text.splitlines():
        ln = ln.rstrip()
        if ln.startswith("## "):
            story.append(Spacer(1, 0.18 * inch))
            story.append(Paragraph(f"<b>{ln[3:].strip()}</b>", styles["Heading2"]))
        elif ln.startswith("### "):
            story.append(Spacer(1, 0.10 * inch))
            story.append(Paragraph(f"<b>{ln[4:].strip()}</b>", styles["Heading3"]))
        elif ln.startswith("- "):
            story.append(Paragraph(f"• {ln[2:].strip()}", styles["Normal"]))
        elif ln.strip() == "":
            story.append(Spacer(1, 0.08 * inch))
        else:
            story.append(
                Paragraph(
                    ln.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"),
                    styles["Normal"],
                )
            )

    doc.build(story, onFirstPage=draw_cover)

def _pick_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    # font umum di ubuntu
    if bold:
        candidates = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]
    else:
        candidates = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        ]
    for fp in candidates:
        try:
            return ImageFont.truetype(fp, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def _text_w(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> int:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0]


def _text_h(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> int:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[3] - bbox[1]


def _wrap_words(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = re.split(r"\s+", (text or "").strip())
    if not words or words == [""]:
        return []
    lines: list[str] = []
    cur = ""
    for w in words:
        test = (cur + " " + w).strip()
        if _text_w(draw, test, font) <= max_width:
            cur = test
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def _fit_text_lines(
    draw: ImageDraw.ImageDraw,
    text: str,
    max_width: int,
    max_lines: int,
    start_size: int,
    min_size: int,
    bold: bool = False,
) -> tuple[list[str], ImageFont.ImageFont]:
    size = start_size
    text = (text or "").strip()

    while size >= min_size:
        font = _pick_font(size, bold=bold)
        lines = _wrap_words(draw, text, font, max_width)
        if lines and len(lines) <= max_lines and all(_text_w(draw, ln, font) <= max_width for ln in lines):
            return lines, font
        size -= 2

    # fallback: pakai min_size dan truncate
    font = _pick_font(min_size, bold=bold)
    lines = _wrap_words(draw, text, font, max_width)
    if not lines:
        return [], font

    lines = lines[:max_lines]
    # truncate last line with ellipsis if still too long
    last = lines[-1]
    ell = "…"
    while last and _text_w(draw, last + ell, font) > max_width:
        last = last[:-1].rstrip()
    lines[-1] = (last + ell) if last else ell
    return lines, font

def _text_width(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> float:
    # textlength tersedia di PIL baru; fallback ke textbbox kalau tidak ada
    try:
        return float(draw.textlength(text, font=font))
    except Exception:
        bbox = draw.textbbox((0, 0), text, font=font)
        return float(bbox[2] - bbox[0])


def _wrap_text_strict(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.ImageFont,
    max_width: int,
    max_lines: int,
) -> list[str]:
    """
    Wrap yang tahan banting:
    - wrap by words
    - kalau ada 1 kata kepanjangan, di-break per karakter
    """
    text = (text or "").strip()
    if not text:
        return []

    words = text.split()
    lines: list[str] = []
    cur = ""

    def flush():
        nonlocal cur
        if cur:
            lines.append(cur)
            cur = ""

    for w in words:
        if not cur:
            test = w
        else:
            test = cur + " " + w

        if _text_width(draw, test, font) <= max_width:
            cur = test
            continue

        # kalau current line ada isi, flush dulu
        if cur:
            flush()

        # sekarang w mungkin terlalu panjang untuk 1 line -> break per karakter
        if _text_width(draw, w, font) <= max_width:
            cur = w
            continue

        chunk = ""
        for ch in w:
            test2 = chunk + ch
            if _text_width(draw, test2, font) <= max_width:
                chunk = test2
            else:
                if chunk:
                    lines.append(chunk)
                chunk = ch
                if len(lines) >= max_lines:
                    break
        if len(lines) >= max_lines:
            cur = ""
            break
        cur = chunk

        if len(lines) >= max_lines:
            break

    if len(lines) < max_lines and cur:
        lines.append(cur)

    # trim to max_lines
    return lines[:max_lines]


def _fit_lines_autoshrink(
    draw: ImageDraw.ImageDraw,
    text: str,
    max_width: int,
    max_lines: int,
    start_size: int,
    min_size: int,
    bold: bool = False,
) -> tuple[list[str], ImageFont.ImageFont]:
    """
    Turunin font sampai:
    - jumlah line <= max_lines
    - semua line <= max_width
    """
    size = start_size
    while size >= min_size:
        font = _pick_font(size, bold=bold)
        lines = _wrap_text_strict(draw, text, font, max_width, max_lines)

        ok = True
        if len(lines) > max_lines:
            ok = False
        else:
            for ln in lines:
                if _text_width(draw, ln, font) > max_width:
                    ok = False
                    break

        if ok:
            return lines, font

        size -= 2

    # fallback: truncate last line
    font = _pick_font(min_size, bold=bold)
    lines = _wrap_text_strict(draw, text, font, max_width, max_lines)
    if not lines:
        return [], font

    lines = lines[:max_lines]
    last = lines[-1]
    ell = "…"
    while last and _text_width(draw, last + ell, font) > max_width:
        last = last[:-1].rstrip()
    lines[-1] = (last + ell) if last else ell
    return lines, font

def _compose_cover(bg: Image.Image, title: str, subtitle: str, author: str) -> Image.Image:
    bg = _to_pil_image(bg)
    img = bg.convert("RGBA")
    W, H = img.size
    draw = ImageDraw.Draw(img)

    # panel teks
    panel_y0 = int(H * 0.08)
    panel_h = int(H * 0.44)
    panel = Image.new("RGBA", (W, panel_h), (0, 0, 0, 0))
    pd = ImageDraw.Draw(panel)
    for y in range(panel_h):
        a = int(175 * (1 - y / max(1, panel_h - 1)))
        pd.rectangle([0, y, W, y + 1], fill=(0, 0, 0, a))
    img.alpha_composite(panel, (0, panel_y0))

    # lebar teks
    margin_x = int(W * 0.07)
    max_width = W - 2 * margin_x

    # fit title/subtitle
    title_lines, title_font = _fit_lines_autoshrink(
        draw,
        title,
        max_width=max_width,
        max_lines=3,
        start_size=max(30, int(H * 0.07)),
        min_size=18,
        bold=True,
    )
    subtitle_lines, sub_font = _fit_lines_autoshrink(
        draw,
        subtitle,
        max_width=max_width,
        max_lines=2,
        start_size=max(18, int(H * 0.034)),
        min_size=12,
        bold=False,
    )

    auth_font = _pick_font(max(14, int(H * 0.028)), bold=False)
    author_line = (author or "").strip()

    # helper gambar center per line
    def draw_center(lines: list[str], font: ImageFont.ImageFont, y: int, fill, line_gap: int) -> int:
        for ln in lines:
            tw = _text_width(draw, ln, font)
            bbox = draw.textbbox((0, 0), ln, font=font)
            th = bbox[3] - bbox[1]
            x = int((W - tw) / 2)
            draw.text((x, y), ln, font=font, fill=fill)
            y += th + line_gap
        return y

    y = panel_y0 + int(H * 0.07)
    y = draw_center(title_lines, title_font, y, fill=(255, 255, 255, 240), line_gap=int(H * 0.012))

    if subtitle_lines:
        y += int(H * 0.01)
        y = draw_center(subtitle_lines, sub_font, y, fill=(255, 255, 255, 220), line_gap=int(H * 0.010))

    if author_line:
        y += int(H * 0.02)
        ln = f"By {author_line}"
        # truncate author kalau kepanjangan
        while ln and _text_width(draw, ln, auth_font) > max_width:
            ln = ln[:-1].rstrip()
        if ln:
            tw = _text_width(draw, ln, auth_font)
            x = int((W - tw) / 2)
            draw.text((x, y), ln, font=auth_font, fill=(255, 255, 255, 210))

    return img.convert("RGB")
    
def _cover_bg_prompt(cfg: dict) -> str:
    title = (cfg.get("title") or "").strip()
    topic = (cfg.get("topic") or "").strip()
    cover = cfg.get("cover") or {}
    style = (cover.get("style") or "Minimal Clean").strip()
    theme = (cover.get("theme") or "").strip()

    # background ONLY: no text
    return (
        "Create an ebook cover BACKGROUND artwork ONLY (no text, no letters, no logos).\n"
        "Leave clear empty space in the upper-middle for title text overlay.\n"
        f"Style: {style}\n"
        f"Book title context: {title}\n"
        f"Topic/brief context: {topic}\n"
        + (f"Theme keywords: {theme}\n" if theme else "")
        + "High quality, clean composition, no watermark.\n"
    )


def _is_transient_error(e: Exception) -> bool:
    msg = str(e)
    name = type(e).__name__
    return (
        "RemoteProtocolError" in name
        or "Server disconnected" in msg
        or "429" in msg
        or "503" in msg
        or "RESOURCE_EXHAUSTED" in msg
        or "UNAVAILABLE" in msg
        or "timeout" in msg.lower()
    )

def _to_pil_image(x) -> Image.Image:
    """
    NanoBananaClient kadang mengembalikan wrapper image (bukan PIL).
    Fungsi ini mengubahnya jadi PIL.Image.Image.
    """
    # sudah PIL
    if hasattr(x, "convert") and hasattr(x, "save"):
        return x

    # wrapper yang punya properti PIL langsung
    for attr in ("pil_image", "pil", "image", "img"):
        v = getattr(x, attr, None)
        if v is not None and hasattr(v, "convert") and hasattr(v, "save"):
            return v

    # wrapper yang menyimpan bytes
    for attr in ("image_bytes", "data", "bytes", "content"):
        b = getattr(x, attr, None)
        if isinstance(b, (bytes, bytearray)) and b:
            return Image.open(io.BytesIO(b)).convert("RGB")

    # kalau x itself bytes
    if isinstance(x, (bytes, bytearray)) and x:
        return Image.open(io.BytesIO(x)).convert("RGB")

    # kalau ada method to_bytes()
    for meth in ("to_bytes", "as_bytes", "get_bytes"):
        m = getattr(x, meth, None)
        if callable(m):
            b = m()
            if isinstance(b, (bytes, bytearray)) and b:
                return Image.open(io.BytesIO(b)).convert("RGB")

    raise TypeError(f"Unsupported image type for cover: {type(x)}")

def _generate_cover_with_retry(
    *,
    api_key: str,
    model: str,
    prompt: str,
    aspect_ratio: str,
    image_size: Optional[str],
    log_path: Path,
    max_attempts: int,
    base_delay: float,
    max_delay: float,
) -> Image.Image:
    client = NanoBananaClient(api_key=api_key, model=model)
    last_err: Exception | None = None

    # Paksa output image (hindari text-only response)
    prompt_img = (
        prompt
        + "\n\nIMPORTANT:\n"
          "- Output MUST be an IMAGE.\n"
          "- Do NOT output any text.\n"
          "- Do NOT include letters/logos.\n"
          "- If you cannot comply, still return an IMAGE background.\n"
    )

    for attempt in range(1, max_attempts + 1):
        try:
            res = client.generate(
                prompt=prompt_img,
                ref_images=None,
                aspect_ratio=aspect_ratio,
                image_size=image_size,
            )

            imgs = list(getattr(res, "images", None) or [])
            if imgs:
                return _to_pil_image(imgs[0])

            # Debug: kadang model balikin text-only (blocked / warning)
            txt = (getattr(res, "text", None) or "").strip()
            if txt:
                _append_log(log_path, f"WARN: cover got TEXT-only response (truncated): {txt[:220]}")

            raise RuntimeError("No images returned from model")

        except Exception as e:
            last_err = e
            msg = str(e)

            # Treat no-image as transient too (biar retry jalan)
            transient = _is_transient_error(e) or ("No images returned" in msg)

            _append_log(log_path, f"WARN: cover attempt {attempt}/{max_attempts} failed: {type(e).__name__}: {e}")

            if not transient:
                break

            delay = min(max_delay, base_delay * (1.6 ** (attempt - 1)))
            time.sleep(delay)

    raise RuntimeError(f"cover generation failed: {type(last_err).__name__}: {last_err}")
    
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", required=True)
    args = ap.parse_args()

    cfg_path = Path(args.config).resolve()
    cfg = json.loads(cfg_path.read_text(encoding="utf-8"))

    job_dir = Path(cfg.get("job_dir") or cfg_path.parent).resolve()
    log_path = Path(cfg.get("log_path") or (job_dir / "job.log")).resolve()

    _append_log(log_path, f"BOOT | cfg={cfg_path} | job_dir={job_dir}")

    api_key = (os.environ.get("GEMINI_API_KEY") or "").strip()
    if not api_key:
        _append_log(log_path, "ERROR: GEMINI_API_KEY env not set.")
        update_progress(job_dir, status="error", total=1, done=0, current="missing GEMINI_API_KEY")
        raise SystemExit(2)

    model = (cfg.get("model") or "gemini-2.5-flash").strip()
    title = (cfg.get("title") or "").strip()
    subtitle = (cfg.get("subtitle") or "").strip()
    author = (cfg.get("author") or "").strip()
    outline_in = (cfg.get("outline") or "").strip()
    chapters_n = int(cfg.get("chapters") or 8)

    # === NEW: output selection ===
    out_cfg = cfg.get("outputs") or {}
    export_md = bool(out_cfg.get("md", True))
    export_docx = bool(out_cfg.get("docx", True))
    export_pdf = bool(out_cfg.get("pdf", True))
    # safety: minimal 1 output
    if not (export_md or export_docx or export_pdf):
        export_md = True

    retry_cfg = cfg.get("retry") or {}
    retry = RetryCfg(
        max_attempts=int(retry_cfg.get("max_attempts") or 6),
        base_delay=float(retry_cfg.get("base_delay") or 1.0),
        max_delay=float(retry_cfg.get("max_delay") or 20.0),
    )

    cover_cfg = cfg.get("cover") or {}
    cover_enabled = bool(cover_cfg.get("enabled", False))
    cover_model = (cover_cfg.get("model") or "gemini-2.5-flash-image").strip()
    cover_aspect = (cover_cfg.get("aspect_ratio") or "3:4").strip()
    cover_size = cover_cfg.get("image_size")  # biasanya None

    # === NEW: progress total dinamis ===
    # outline(1) + chapters(N) + finalize(assemble=1 + docx? + pdf?)
    cover_step = 1 if cover_enabled else 0
    final_steps = 1 + (1 if export_docx else 0) + (1 if export_pdf else 0)  # assemble + docx? + pdf?
    total = max(1, 1 + cover_step + chapters_n + final_steps)

    init_progress(job_dir, total)
    done = 0

    out_dir = (job_dir / "outputs").resolve()
    meta_dir = (job_dir / "meta").resolve()
    ch_dir = (job_dir / "chapters").resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    meta_dir.mkdir(parents=True, exist_ok=True)
    ch_dir.mkdir(parents=True, exist_ok=True)

    _append_log(
        log_path,
        f"JOB START | model={model} chapters={chapters_n} outputs=md:{export_md},docx:{export_docx},pdf:{export_pdf}",
    )

    # 1) Outline
    if outline_in:
        _append_log(log_path, "Using user-provided outline.")
        chapters = _parse_outline(outline_in, chapters_n)
    else:
        update_progress(job_dir, status="running", total=total, done=done, current="Generating outline")
        _append_log(log_path, "Generating outline...")
        outline_txt = _gen_text(
            api_key=api_key,
            model=model,
            prompt=_outline_prompt(cfg),
            log_path=log_path,
            retry=retry,
            temperature=0.4,
        )
        (meta_dir / "outline_raw.txt").write_text(outline_txt, encoding="utf-8")
        chapters = _parse_outline(outline_txt, chapters_n)

    (meta_dir / "outline.json").write_text(
        json.dumps({"chapters": chapters}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    done += 1
    update_progress(job_dir, status="running", total=total, done=done, current="Outline ready")

    cover_path: Optional[Path] = None
    if cover_enabled:
        current = "Generating cover"
        update_progress(job_dir, status="running", total=total, done=done, current=current)
        _append_log(log_path, current)

        try:
            bg = _generate_cover_with_retry(
                api_key=api_key,
                model=cover_model,
                prompt=_cover_bg_prompt(cfg),
                aspect_ratio=cover_aspect,
                image_size=cover_size,
                log_path=log_path,
                max_attempts=retry.max_attempts,
                base_delay=retry.base_delay,
                max_delay=retry.max_delay,
            )
            # compose overlay text
            cover_img = _compose_cover(bg, title=title, subtitle=subtitle, author=author)

            cover_path = (out_dir / "cover.png")
            cover_img.save(cover_path)
            _append_log(log_path, f"Cover saved: {cover_path.name}")
        except Exception as e:
            _append_log(log_path, f"WARN: cover generation failed: {type(e).__name__}: {e}")

        done += 1
        update_progress(job_dir, status="running", total=total, done=done, current="Cover ready")

    # 2) Chapters
    md_parts: List[str] = []
    if title:
        md_parts.append(f"# {title}")
    if subtitle:
        md_parts.append(f"**{subtitle}**")
    if author:
        md_parts.append(f"_By {author}_")
    md_parts.append("\n---\n")
    md_parts.append("## Table of Contents")
    for i, ch in enumerate(chapters, start=1):
        md_parts.append(f"- Chapter {i}: {ch['title']}")
    md_parts.append("\n---\n")

    for i, ch in enumerate(chapters, start=1):
        current = f"Writing Chapter {i}/{chapters_n}: {ch['title']}"
        update_progress(job_dir, status="running", total=total, done=done, current=current)
        _append_log(log_path, current)

        chapter_md = _gen_text(
            api_key=api_key,
            model=model,
            prompt=_chapter_prompt(cfg, i, ch),
            log_path=log_path,
            retry=retry,
            temperature=0.75,
        )

        fn = f"{i:02d}_{_safe_slug(ch['title'])}.md"
        (ch_dir / fn).write_text(chapter_md, encoding="utf-8")
        md_parts.append(chapter_md)
        md_parts.append("\n")

        done += 1
        update_progress(job_dir, status="running", total=total, done=done, current=current)

    # 3) Finalize outputs (sesuai pilihan)
    update_progress(job_dir, status="running", total=total, done=done, current="Assembling outputs")
    full_md = "\n\n".join(md_parts).strip() + "\n"

    # Markdown optional
    if export_md:
        (out_dir / "book.md").write_text(full_md, encoding="utf-8")
        _append_log(log_path, "Markdown saved.")
    else:
        _append_log(log_path, "Markdown skipped (not selected).")

    done += 1
    update_progress(job_dir, status="running", total=total, done=done, current="Writing DOCX/PDF")

    # DOCX optional
    if export_docx:
        try:
            _write_docx(full_md, out_dir / "book.docx", title, subtitle, author, cover_path=cover_path)
            _append_log(log_path, "DOCX saved.")
        except Exception as e:
            _append_log(log_path, f"WARN: DOCX failed: {type(e).__name__}: {e}")
        done += 1
        update_progress(job_dir, status="running", total=total, done=done, current="DOCX done")
    else:
        _append_log(log_path, "DOCX skipped (not selected).")

    # PDF optional
    if export_pdf:
        try:
            _write_pdf(full_md, out_dir / "book.pdf", title, subtitle, author, cover_path=cover_path)
            _append_log(log_path, "PDF saved.")
        except Exception as e:
            _append_log(log_path, f"WARN: PDF failed: {type(e).__name__}: {e}")
        done += 1
        update_progress(job_dir, status="running", total=total, done=done, current="PDF done")
    else:
        _append_log(log_path, "PDF skipped (not selected).")

    update_progress(job_dir, status="done", total=total, done=total, current="")
    _append_log(log_path, "JOB DONE")

if __name__ == "__main__":
    main()
